# faculty_duty_app.py
import streamlit as st
import pandas as pd
import datetime
from collections import defaultdict
from io import BytesIO
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.table import WD_ROW_HEIGHT_RULE
import importlib
import os
import pickle
import collections

st.set_page_config(page_title="Faculty Duty Assignment System", layout="wide")

# --- Utility Functions ---
def to_ddmmyyyy(date):
    if pd.isnull(date) or date is None:
        return ''
    if isinstance(date, str):
        try:
            date = pd.to_datetime(date, errors='coerce').date()
            if pd.isnull(date):
                return ''
        except Exception:
            return ''
    if isinstance(date, (datetime.datetime, pd.Timestamp)):
        if pd.isnull(date):
            return ''
        date = date.date()
    if isinstance(date, datetime.date):
        return date.strftime('%d-%m-%Y')
    return ''

def from_ddmmyyyy(date_str):
    try:
        return pd.to_datetime(date_str, format='%d-%m-%Y', errors='coerce').date()
    except Exception:
        return None

def normalize_dates(df, col='Date'):
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], errors='coerce').dt.date
    return df

def normalize_schedule(schedule):
    for day in schedule:
        if not isinstance(day['date'], (pd.Timestamp, )):
            day['date'] = pd.to_datetime(day['date'], errors='coerce')
        if hasattr(day['date'], 'date') and day['date'] is not None and not pd.isnull(day['date']):
            day['date'] = day['date'].date()
    return schedule

def normalize_unavailability(unavailability):
    for faculty in unavailability:
        for half in ['first_half', 'second_half']:
            new_dates = set()
            for d in unavailability[faculty][half]:
                try:
                    parsed_date = pd.to_datetime(d, errors='coerce').date()
                    if parsed_date is not None and not pd.isnull(parsed_date):
                        new_dates.add(parsed_date)
                except Exception:
                    pass
            unavailability[faculty][half] = new_dates
    return unavailability

# --- Constraint Validation ---
def validate_assignment_constraints(df, faculty_list, max_duties_dict, unavailability, faculty_groups, exam_schedule):
    errors = []
    is_valid = True
    # 1. Faculty unavailability
    for _, row in df.iterrows():
        faculty = row['Faculty']
        date = row['Date']
        shift = row['Shift']
        shift_type = 'first_half' if shift == 'First Half' else 'second_half'
        if faculty in unavailability:
            unavail_dates = set()
            for d in unavailability[faculty][shift_type]:
                try:
                    unavail_dates.add(pd.to_datetime(d, errors='coerce').date())
                except Exception:
                    pass
            if date in unavail_dates:
                is_valid = False
                errors.append(f"{faculty} is assigned on {to_ddmmyyyy(date)} {shift} but marked as unavailable.")
    # 2. Max duties
    duty_counts = defaultdict(int)
    for _, row in df.iterrows():
        faculty = row['Faculty']
        duty_counts[faculty] += 1
    for faculty in faculty_list:
        max_allowed = max_duties_dict.get(faculty, float('inf'))
        if duty_counts[faculty] > max_allowed:
            is_valid = False
            errors.append(f"{faculty} has {duty_counts[faculty]} duties assigned, exceeding maximum of {max_allowed}.")
    # 3. Group togetherness
    for group in faculty_groups:
        for _, row in df.iterrows():
            date = row['Date']
            shift = row['Shift']
            assigned = df[(df['Date'] == date) & (df['Shift'] == shift)]['Faculty'].tolist()
            if any(f in assigned for f in group):
                if not all(f in assigned for f in group):
                    is_valid = False
                    errors.append(f"Group {', '.join(group)} not assigned together on {to_ddmmyyyy(date)} {shift}.")
    # 4. Required number of faculty per shift
    for day in exam_schedule:
        sched_date = day['date']
        for shift, label in [('First Half', 'first_half'), ('Second Half', 'second_half')]:
            required = day[label] if label in day else 0
            assigned = len(df[(df['Date'] == sched_date) & (df['Shift'] == shift)])
            if assigned != required:
                is_valid = False
                errors.append(f"{assigned} faculty assigned on {to_ddmmyyyy(sched_date)} {shift}, required: {required}.")
    # 5. No faculty assigned to both shifts on the same day
    same_day_double = set()
    for faculty in faculty_list:
        dates = df[df['Faculty'] == faculty]['Date']
        date_counts = dates.value_counts()
        for d, count in date_counts.items():
            if count > 1:
                is_valid = False
                errors.append(f"{faculty} is assigned to both shifts on {to_ddmmyyyy(d)}.")
                same_day_double.add((faculty, d))
    return is_valid, errors, list(same_day_double)

# --- Report Generation Functions ---

def generate_faculty_summary_excel(df, unavailability=None):
    try:
        faculty_list = None
        if hasattr(st.session_state, 'faculty_df') and st.session_state.faculty_df is not None:
            faculty_list = st.session_state.faculty_df['faculty'].tolist()
        required_columns = {'Faculty', 'Date', 'Shift'}
        if df is None or df.empty or (set(df.columns) & required_columns) != required_columns:
            df = pd.DataFrame({col: pd.Series(dtype='object') for col in ['Faculty', 'Date', 'Shift']})
        else:
            df = df.copy()
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
        if faculty_list is not None:
            summary_rows = []
            for faculty in faculty_list:
                faculty_df = df[df['Faculty'] == faculty]
                first_half_count = (faculty_df['Shift'] == 'First Half').sum()
                second_half_count = (faculty_df['Shift'] == 'Second Half').sum()
                total_duties = len(faculty_df)
                fh_list = faculty_df.loc[faculty_df['Shift'] == 'First Half', 'Date'].to_list()
                sh_list = faculty_df.loc[faculty_df['Shift'] == 'Second Half', 'Date'].to_list()
                fh_dates = pd.Series(list(pd.to_datetime(fh_list, errors='coerce')), dtype='datetime64[ns]')
                sh_dates = pd.Series(list(pd.to_datetime(sh_list, errors='coerce')), dtype='datetime64[ns]')
                if not fh_dates.empty and pd.api.types.is_datetime64_any_dtype(fh_dates):
                    first_half_dates = ', '.join(fh_dates.dt.strftime('%d-%m-%Y'))
                else:
                    first_half_dates = ''
                if not sh_dates.empty and pd.api.types.is_datetime64_any_dtype(sh_dates):
                    second_half_dates = ', '.join(sh_dates.dt.strftime('%d-%m-%Y'))
                else:
                    second_half_dates = ''
                if unavailability:
                    fh_unavail = ', '.join(d.strftime('%d-%m-%Y') for d in sorted(unavailability.get(faculty, {'first_half': set()})['first_half'])) or 'None'
                    sh_unavail = ', '.join(d.strftime('%d-%m-%Y') for d in sorted(unavailability.get(faculty, {'second_half': set()})['second_half'])) or 'None'
                    total_unavail = len(unavailability.get(faculty, {'first_half': set(), 'second_half': set()})['first_half']) + \
                                   len(unavailability.get(faculty, {'first_half': set(), 'second_half': set()})['second_half'])
                else:
                    fh_unavail = 'None'
                    sh_unavail = 'None'
                    total_unavail = 0
                summary_rows.append({
                    'Faculty': faculty,
                    'First Half Duties': first_half_count,
                    'Second Half Duties': second_half_count,
                    'Total Duties': total_duties,
                    'First Half Dates': first_half_dates,
                    'Second Half Dates': second_half_dates,
                    'First Half Unavailable': fh_unavail,
                    'Second Half Unavailable': sh_unavail,
                    'Total Unavailable Slots': total_unavail
                })
            faculty_summary = pd.DataFrame(summary_rows)
        else:
            faculty_summary = df.groupby('Faculty').agg({
                'Shift': lambda x: [
                    sum(x == 'First Half'),
                    sum(x == 'Second Half'),
                    len(x)
                ]
            }).reset_index()
            faculty_summary[['First Half Duties', 'Second Half Duties', 'Total Duties']] = pd.DataFrame(
                faculty_summary['Shift'].tolist(), 
                index=faculty_summary.index
            )
            faculty_summary.drop('Shift', axis=1, inplace=True)
            def get_shift_dates(faculty, shift):
                dates = df[(df['Faculty'] == faculty) & (df['Shift'] == shift)]['Date']
                if not isinstance(dates, pd.Series):
                    dates = pd.Series(dates)
                if not pd.api.types.is_datetime64_any_dtype(dates):
                    dates = pd.to_datetime(dates, errors='coerce')
                dates = dates.dropna()
                if not dates.empty:
                    return ', '.join(dates.dt.strftime('%d-%m-%Y'))
                else:
                    return ''
            faculty_summary['First Half Dates'] = faculty_summary['Faculty'].apply(
                lambda f: get_shift_dates(f, 'First Half')
            )
            faculty_summary['Second Half Dates'] = faculty_summary['Faculty'].apply(
                lambda f: get_shift_dates(f, 'Second Half')
            )
            if unavailability:
                faculty_summary['First Half Unavailable'] = faculty_summary['Faculty'].apply(
                    lambda f: ', '.join(d.strftime('%d-%m-%Y') for d in sorted(unavailability.get(f, {'first_half': set()})['first_half'])) or 'None'
                )
                faculty_summary['Second Half Unavailable'] = faculty_summary['Faculty'].apply(
                    lambda f: ', '.join(d.strftime('%d-%m-%Y') for d in sorted(unavailability.get(f, {'second_half': set()})['second_half'])) or 'None'
                )
                faculty_summary['Total Unavailable Slots'] = faculty_summary['Faculty'].apply(
                    lambda f: len(unavailability.get(f, {'first_half': set(), 'second_half': set()})['first_half']) +
                            len(unavailability.get(f, {'first_half': set(), 'second_half': set()})['second_half'])
                )
            else:
                faculty_summary['First Half Unavailable'] = 'None'
                faculty_summary['Second Half Unavailable'] = 'None'
                faculty_summary['Total Unavailable Slots'] = 0
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            faculty_summary.to_excel(writer, sheet_name="Faculty Duty Summary", index=False)
            worksheet = writer.sheets["Faculty Duty Summary"]
            for idx, col in enumerate(faculty_summary.columns):
                max_length = max(
                    faculty_summary[col].astype(str).apply(len).max(),
                    len(col)
                ) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error generating Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def generate_word_doc(df):
    df = df.copy()
    # Always keep Date as datetime.date for logic, only format for display
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce").dt.date
    # Drop rows where Date is missing
    df = df[df["Date"].notna()]
    doc = Document()
    # Set narrow margins (0.5 inches = 12.7 mm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.left_margin = Mm(12.7)
        section.right_margin = Mm(12.7)
    # Get exam type, semester and year from session state
    exam_type = st.session_state.get('exam_type', 'MID SEM')  # Default to MID SEM
    semester = st.session_state.get('semester', 'MO')  # Default to MO
    year = st.session_state.get('year', '2025')  # Default to current year
    # Add current date at top right
    today_str = datetime.datetime.today().strftime('%d/%m/%Y')
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Date : {today_str}')
    date_para.alignment = 2  # Right align
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(12)
    # Create header with exam type, semester and year
    p = doc.add_paragraph()
    run1 = p.add_run("Department of Computer Science & Engineering")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run1.add_break()
    run2 = p.add_run("BIT MESRA , RANCHI")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    p.alignment = 1  # Center align (optional)
    header = f"Examination Duty Chart - {exam_type} {semester} {year}"
    p2 = doc.add_paragraph()
    run_header = p2.add_run(header)
    run_header.font.name = 'Times New Roman'
    run_header.font.size = Pt(14)
    p2.alignment = 1  # Center align (optional)
    # Add underline to header
    run_header.font.underline = True
    # Add time paragraph based on exam_type
    if exam_type == "MID SEM":
        time_text = (
            "Time: 09.40 A.M. to 12.00 NOON (1st Half)\n"
            "01.40 P.M. to 04.00 P.M. (2nd Half)"
        )
    else:
        time_text = (
            "Time: 09.40 A.M. to 01.00 P.M. (1st Half)\n"
            "01.40 P.M. to 05.00 P.M. (2nd Half)"
        )
    p3 = doc.add_paragraph()
    run_time = p3.add_run(time_text)
    run_time.font.name = 'Times New Roman'
    run_time.font.size = Pt(12)
    p3.alignment = 1  # Center align (optional)
    # Get unique dates and sort them
    unique_dates = sorted(df["Date"].unique())
    for date in unique_dates:
        if pd.isna(date):
            continue
        # Add date display before the table as bold paragraph
        doc.add_paragraph()  # Add some space before the date display
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(date.strftime("%d-%m-%Y"))
        date_run.bold = True
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        # Filter DataFrame for the current date
        df_for_date = df[df["Date"] == date].copy()
        # Sort by shift for correct merging order
        df_for_date = df_for_date.sort_values(by=["Shift"])
        # Calculate total rows needed including blank row between shifts
        total_rows = len(df_for_date) + 1  # +1 for header
        if "First Half" in df_for_date["Shift"].values and "Second Half" in df_for_date["Shift"].values:
            total_rows += 1  # Add one more row for blank row between shifts
        # Create table with calculated rows
        table = doc.add_table(rows=total_rows, cols=5)
        table.style = "Table Grid"
        # Set column widths (in mm)
        table.columns[0].width = Mm(25)  # Shift column
        table.columns[1].width = Mm(15.1)  # S.No column (1.51 cm = 15.1 mm)
        table.columns[2].width = Mm(60)  # Faculty column
        table.columns[3].width = Mm(30)  # Phone No column
        table.columns[4].width = Mm(40)  # Email ID column
        # Add header row
        hdr_cells = table.rows[0].cells
        headers = ["Shift", "S.No", "Faculty", "Phone No", "Email ID"]
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header
            # Center align the text horizontally
            cell.paragraphs[0].alignment = 1
            # Make text bold
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        # Add data rows for this date
        current_row = 1  # Start from row 1 (after header)
        serial_no = 1  # Reset serial number for each date's table
        first_half_end_row = None  # Track where First Half ends
        # First pass: Process all First Half rows
        for idx, row in df_for_date.iterrows():
            if row["Shift"] == "First Half":
                while current_row >= len(table.rows):
                    table.add_row()
                row_cells = table.rows[current_row].cells
                row_cells[0].text = str(row["Shift"]) if current_row == 1 else ""
                row_cells[1].text = str(serial_no)
                row_cells[2].text = str(row["Faculty"])
                row_cells[3].text = str(row.get("Phone No", ""))
                row_cells[4].text = str(row.get("Email Id", ""))
                for cell in row_cells:
                    cell.paragraphs[0].alignment = 1
                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                serial_no += 1
                current_row += 1
                first_half_end_row = current_row - 1
        # Add blank row after First Half if Second Half exists
        if first_half_end_row is not None and "Second Half" in df_for_date["Shift"].values:
            while current_row >= len(table.rows):
                table.add_row()
            blank_cells = table.rows[current_row].cells
            for cell in blank_cells:
                cell.text = ""
            current_row += 1
        # Second pass: Process all Second Half rows
        serial_no = 1
        second_half_start_row = current_row
        for idx, row in df_for_date.iterrows():
            if row["Shift"] == "Second Half":
                while current_row >= len(table.rows):
                    table.add_row()
                row_cells = table.rows[current_row].cells
                row_cells[0].text = str(row["Shift"]) if current_row == second_half_start_row else ""
                row_cells[1].text = str(serial_no)
                row_cells[2].text = str(row["Faculty"])
                row_cells[3].text = str(row.get("Phone No", ""))
                row_cells[4].text = str(row.get("Email Id", ""))
                for cell in row_cells:
                    cell.paragraphs[0].alignment = 1
                    for run in cell.paragraphs[0].runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                serial_no += 1
                current_row += 1
        # Merge cells for First Half
        if first_half_end_row is not None:
            try:
                merged_cell = table.cell(1, 0).merge(table.cell(first_half_end_row, 0))
                merged_cell.vertical_alignment = WD_ROW_HEIGHT_RULE.AT_LEAST
            except Exception as e:
                print(f"Error merging First Half cells: {e}")
        # Merge cells for Second Half
        if "Second Half" in df_for_date["Shift"].values:
            try:
                merged_cell = table.cell(second_half_start_row, 0).merge(table.cell(current_row - 1, 0))
                merged_cell.vertical_alignment = WD_ROW_HEIGHT_RULE.AT_LEAST
            except Exception as e:
                print(f"Error merging Second Half cells: {e}")
    # Add a note section at the end
    doc.add_paragraph()
    doc.add_heading("Note:", level=1)
    notes = [
        "All the Invigilators according to the invigilation chart are requested to report to the upstairs examination office 20 minute before the examination starts (The room allotment will be done before the start of each examination).",
        "If any Invigilator is unable to do invigilation duty for any reason, then it should be brought to the notice of the Controller of Examination with alternative arrangement through HoD well before the start of the examination.",
        "Invigilators will be prohibited from carrying and using cell phones in the Examination Hall (As recommended in the 66th meeting of the Examination Committee meeting).",
        "Invigilators should make sure that bags of the students are not kept inside the Examination Hall (As recommended in the 66th meeting of the Examination Committee meeting)."
    ]
    for i, note in enumerate(notes, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}.      ").bold = True
        p.add_run(note)
    # Add signature section
    doc.add_paragraph("\n\n")
    signature = doc.add_paragraph()
    signature.add_run("(Dr. A. Mustafi)\n").bold = True
    signature.add_run("Professor & Head\n")
    signature.add_run("Department of Computer Science & Engineering\n")
    signature.add_run("B.I.T., Mesra, Ranchi")
    # Add copy to section
    doc.add_paragraph("\n")
    copy_to = [
        "All faculty members (through email)",
        "Controller of examination", 
        "Copy to V.C Office",
        "Office File"
    ]
    p_heading = doc.add_paragraph()
    p_heading.add_run("Copy to:").bold = True
    p_recipients = doc.add_paragraph()
    for i, recipient in enumerate(copy_to, 1):
        run = p_recipients.add_run(f"{i}.\t{recipient}")
        if i < len(copy_to):
            run.add_break() # Add a line break instead of a new paragraph
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- Auto-load and Resume/Start Fresh Logic ---
JOB_FILE = "last_job.pkl"

def save_job():
    # Convert faculty_unavailability to a regular dict for pickling
    faculty_unavailability = st.session_state.get('faculty_unavailability')
    if isinstance(faculty_unavailability, collections.defaultdict):
        faculty_unavailability = dict(faculty_unavailability)
    job_state = {
        'faculty_df': st.session_state.get('faculty_df'),
        'faculty_unavailability': faculty_unavailability,
        'assigned_duty_df': st.session_state.get('assigned_duty_df'),
        'faculty_groups': st.session_state.get('faculty_groups'),
        'max_duties_dict': st.session_state.get('max_duties_dict'),
        'exam_schedule': st.session_state.get('exam_schedule'),
        'manual_selected': st.session_state.get('manual_selected'),
        'exam_type': st.session_state.get('exam_type', 'MID SEM'),
        'semester': st.session_state.get('semester', 'MO'),
        'year': st.session_state.get('year', '2025'),
    }
    with open(JOB_FILE, 'wb') as f:
        pickle.dump(job_state, f)

def load_job():
    if os.path.exists(JOB_FILE):
        with open(JOB_FILE, 'rb') as f:
            job_state = pickle.load(f)
        for k, v in job_state.items():
            if k == 'faculty_unavailability':
                from collections import defaultdict
                d = v if v is not None else {}
                st.session_state[k] = defaultdict(lambda: {"first_half": set(), "second_half": set()}, d)
            else:
                st.session_state[k] = v
        return True
    return False

def clear_job():
    if os.path.exists(JOB_FILE):
        os.remove(JOB_FILE)
    for k in ['faculty_df', 'faculty_unavailability', 'assigned_duty_df', 'faculty_groups', 'max_duties_dict', 'exam_schedule', 'manual_selected', 'exam_type', 'semester', 'year']:
        if k in st.session_state:
            del st.session_state[k]

# --- Prompt to Resume or Start Fresh ---
if 'job_prompt_done' not in st.session_state:
    if os.path.exists(JOB_FILE):
        st.markdown("""
        <div style='background-color:#f8f9fa; padding:1rem; border-radius:8px; border:1px solid #ddd;'>
        <b>Resume your last job?</b><br>
        A previous session was found. Would you like to resume where you left off or start fresh?
        </div>
        """, unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Resume last job"):
                load_job()
                st.session_state['job_prompt_done'] = True
                st.rerun()
        with col2:
            if st.button("Start fresh"):
                clear_job()
                st.session_state['job_prompt_done'] = True
                st.rerun()
        st.stop()
    else:
        st.session_state['job_prompt_done'] = True

# --- After any major change, auto-save the job ---
def auto_save_job():
    save_job()

# Call auto_save_job after any major change (faculty, exam dates, assignments, etc.)
# Example: after uploading faculty, exam schedule, or generating assignments, call auto_save_job()

# --- Streamlit App ---

# --- Sidebar: Configuration ---
# Remove exam_type, semester, year from sidebar
# Add custom CSS for selectbox label and selected value
st.markdown(
    '''
    <style>
    /* Style the sidebar label */
    .css-1v0mbdj, .css-1cpxqw2 {
        color: #1a237e !important;
        font-weight: bold !important;
        font-size: 1.1rem !important;
    }
    /* Style the selected value in selectbox */
    .stSelectbox > div > div {
        color: #1a237e !important;
        font-weight: 600 !important;
    }
    </style>
    ''',
    unsafe_allow_html=True
)
st.sidebar.header(" MAIN SECTIONS")
# Remove exam_type, semester, year from sidebar
# Add configuration to main area at the top
st.markdown("## 🛠️ Configuration")
col_exam, col_sem, col_year = st.columns([2,2,1])
with col_exam:
    exam_type_options = ["MID SEM", "END SEM"]
    exam_type = st.selectbox("Exam Type", exam_type_options, index=exam_type_options.index(st.session_state.get('exam_type', 'MID SEM')) if st.session_state.get('exam_type', 'MID SEM') in exam_type_options else 0, key="main_exam_type")
with col_sem:
    semester_options = ["MO", "SP"]
    semester = st.selectbox("Semester", semester_options, index=semester_options.index(st.session_state.get('semester', 'MO')) if st.session_state.get('semester', 'MO') in semester_options else 0, key="main_semester")
with col_year:
    year = st.text_input("Year", "2025", key="main_year")

# --- Session State Initialization ---
if 'faculty_df' not in st.session_state:
    st.session_state.faculty_df = None
if 'faculty_unavailability' not in st.session_state:
    st.session_state.faculty_unavailability = defaultdict(lambda: {"first_half": set(), "second_half": set()})
if 'assigned_duty_df' not in st.session_state:
    st.session_state.assigned_duty_df = None
if 'faculty_groups' not in st.session_state:
    st.session_state.faculty_groups = []
if 'max_duties_dict' not in st.session_state:
    st.session_state.max_duties_dict = {}
if 'exam_schedule' not in st.session_state:
    st.session_state.exam_schedule = []

# --- Main Navigation ---
section = st.sidebar.radio(
    "Navigate",
    [
        "Faculty Management",
        "Exam Schedule",
        "Duty Assignment",
        "Manual Intervention",
        "Reports & Downloads"
    ]
)

# --- Faculty Management ---
if section == "Faculty Management":
    st.header("Faculty Management")
    uploaded_file = st.file_uploader("Upload Faculty Details (CSV or Excel)", type=["csv", "xlsx"])
    if uploaded_file is not None:
        if uploaded_file.name.endswith('.csv'):
            faculty_df = pd.read_csv(uploaded_file)
        else:
            faculty_df = pd.read_excel(uploaded_file)
        # Clean and validate
        faculty_df['faculty'] = faculty_df['faculty'].astype(str).str.strip()
        missing_names = faculty_df['faculty'].isna() | (faculty_df['faculty'] == '') | (faculty_df['faculty'].str.lower() == 'nan')
        num_missing = missing_names.sum()
        if num_missing > 0:
            st.warning(f"⚠️ {num_missing} row(s) with missing faculty names were removed.")
            faculty_df = faculty_df[~missing_names]
        duplicate_names = faculty_df['faculty'].duplicated(keep='first')
        if duplicate_names.any():
            dups = faculty_df.loc[duplicate_names, 'faculty'].tolist()
            st.warning(f"⚠️ Duplicate faculty names removed: {', '.join(dups)}")
            faculty_df = faculty_df[~duplicate_names]
        faculty_df = faculty_df.reset_index(drop=True)
        st.session_state.faculty_df = faculty_df
        # Set max duties dict from file or default
        if 'Max Duties' in faculty_df.columns:
            st.session_state.max_duties_dict = dict(zip(faculty_df['faculty'], faculty_df['Max Duties']))
        else:
            st.session_state.max_duties_dict = {f: 3 for f in faculty_df['faculty']}
            st.warning("No 'Max Duties' column found in the uploaded file. Defaulting all to 3. Please add this column for per-faculty limits.")
        # Ensure unavailability entry for every faculty
        for faculty in faculty_df['faculty']:
            if faculty not in st.session_state.faculty_unavailability:
                st.session_state.faculty_unavailability[faculty] = {"first_half": set(), "second_half": set()}
        st.success("✅ Faculty details uploaded and cleaned!")
        st.dataframe(faculty_df)
        auto_save_job()
    # Always show current faculty_df if available
    elif st.session_state.faculty_df is not None:
        st.info("Current faculty list loaded:")
        st.dataframe(st.session_state.faculty_df)
    else:
        st.info("Please upload a faculty details file to get started.")
    # Grouping, max duties, and unavailability UI can be added here...
    st.markdown('---')
    st.subheader('Faculty Grouping (Assign Together)')
    if 'faculty_groups' not in st.session_state:
        st.session_state.faculty_groups = []
    faculty_list = st.session_state.faculty_df['faculty'].tolist() if st.session_state.faculty_df is not None else []
    # Show current groups
    if st.session_state.faculty_groups:
        st.markdown('**Current Groups:**')
        for idx, group in enumerate(st.session_state.faculty_groups, 1):
            st.markdown(f"- Group {idx}: {', '.join(group)}")
    else:
        st.info('No groups defined yet.')
    # Add new group
    st.markdown('**Create a New Group (2 or more faculty):**')
    new_group = st.multiselect('Select faculty to group together', options=faculty_list, key='new_group_select')
    if st.button('Add Group'):
        if len(new_group) < 2:
            st.warning('Please select at least 2 faculty for a group.')
        elif any(set(new_group) == set(g) for g in st.session_state.faculty_groups):
            st.warning('This group already exists.')
        else:
            st.session_state.faculty_groups.append(list(new_group))
            st.success(f"Group added: {', '.join(new_group)}")
            st.rerun()
    # Option to delete groups
    if st.session_state.faculty_groups:
        group_to_delete = st.selectbox('Delete a group', options=[f'Group {i+1}: {", ".join(g)}' for i, g in enumerate(st.session_state.faculty_groups)], key='delete_group_select')
        if st.button('Delete Selected Group'):
            idx = int(group_to_delete.split(':')[0].replace('Group ', '')) - 1
            st.session_state.faculty_groups.pop(idx)
            st.success('Group deleted.')
            st.rerun()
    st.markdown('---')
    st.subheader('Faculty Unavailability (by Date & Shift)')
    faculty_search = st.text_input('Search Faculty', '', key='faculty_unavail_search')
    st.caption('Type and press Enter or click outside the box to filter.')
    # Clean faculty list: remove NaN, empty, and strip spaces, ensure all are strings
    faculty_list = st.session_state.faculty_df['faculty'].tolist() if st.session_state.faculty_df is not None else []
    faculty_list = [str(f).strip() for f in faculty_list if pd.notna(f) and str(f).strip() and str(f).strip().lower() != 'nan']
    # Filter by search (case and space insensitive)
    search_val = faculty_search.strip().lower().replace(' ', '')
    if search_val:
        faculty_list = [f for f in faculty_list if search_val in f.lower().replace(' ', '')]
    if not faculty_list:
        st.warning('No faculty found matching your search.')
    if not st.session_state.exam_schedule:
        st.warning("Please configure the exam schedule first!")
    else:
        for faculty in faculty_list:
            st.markdown(f"**{faculty}**")
            cols = st.columns(len(st.session_state.exam_schedule) * 2)
            for i, date in enumerate(st.session_state.exam_schedule):
                with cols[2*i]:
                    key_fh = f"unavail_{faculty}_fh_{date['date']}"
                    unavailable_fh = st.checkbox(
                        f"{to_ddmmyyyy(date['date'])} First Half",
                        value=date['date'] in st.session_state.faculty_unavailability[faculty]['first_half'],
                        key=key_fh
                    )
                    if unavailable_fh:
                        st.session_state.faculty_unavailability[faculty]['first_half'].add(date['date'])
                    else:
                        st.session_state.faculty_unavailability[faculty]['first_half'].discard(date['date'])
                with cols[2*i+1]:
                    key_sh = f"unavail_{faculty}_sh_{date['date']}"
                    unavailable_sh = st.checkbox(
                        f"{to_ddmmyyyy(date['date'])} Second Half",
                        value=date['date'] in st.session_state.faculty_unavailability[faculty]['second_half'],
                        key=key_sh
                    )
                    if unavailable_sh:
                        st.session_state.faculty_unavailability[faculty]['second_half'].add(date['date'])
                    else:
                        st.session_state.faculty_unavailability[faculty]['second_half'].discard(date['date'])

# --- Exam Schedule ---
elif section == "Exam Schedule":
    st.header("Exam Schedule")
    # Always show current exam schedule if available
    if st.session_state.exam_schedule:
        st.dataframe(pd.DataFrame({
            'Date': [to_ddmmyyyy(day['date']) for day in st.session_state.exam_schedule],
            'First Half Required': [day['first_half'] for day in st.session_state.exam_schedule],
            'Second Half Required': [day['second_half'] for day in st.session_state.exam_schedule]
        }))
    num_exam_dates = st.number_input("Number of Exam Days Needed", min_value=1, max_value=30, value=len(st.session_state.exam_schedule) or 5)
    temp_schedule = []
    for idx in range(num_exam_dates):
        st.markdown(f"#### Exam Day {idx + 1}")
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            selected_date = st.date_input("Date", key=f"exam_date_{idx}", value=st.session_state.exam_schedule[idx]['date'] if idx < len(st.session_state.exam_schedule) else None)
            if selected_date is not None:
                st.markdown(f"<span style='font-size:1.3rem; font-weight:600;'>Selected: {selected_date.strftime('%d-%m-%Y')}</span>", unsafe_allow_html=True)
            else:
                st.markdown(f"<span style='font-size:1.3rem; font-weight:600; color: #d9534f;'>Please select a date</span>", unsafe_allow_html=True)
        with col2:
            first_half_count = st.number_input("First Half Faculty", min_value=0, value=st.session_state.exam_schedule[idx]['first_half'] if idx < len(st.session_state.exam_schedule) else 2, key=f"first_half_{idx}")
        with col3:
            second_half_count = st.number_input("Second Half Faculty", min_value=0, value=st.session_state.exam_schedule[idx]['second_half'] if idx < len(st.session_state.exam_schedule) else 2, key=f"second_half_{idx}")
        temp_schedule.append({
            'date': selected_date,
            'first_half': first_half_count,
            'second_half': second_half_count
        })
        st.markdown("<div style='margin-bottom: -1.5rem'></div>", unsafe_allow_html=True)
    if st.button("✅ Confirm Exam Schedule"):
        st.session_state.exam_schedule = normalize_schedule(temp_schedule)
        st.success("✅ Successfully configured exam schedule!")
        auto_save_job()

# --- Duty Assignment ---
elif section == "Duty Assignment":
    st.header("Duty Assignment")
    # Show current exam schedule
    if not st.session_state.exam_schedule:
        st.warning("Please configure the exam schedule first.")
        st.stop()
    if st.session_state.faculty_df is None or st.session_state.faculty_df.empty:
        st.warning("Please upload faculty details first.")
        st.stop()
    st.subheader("Current Exam Schedule")
    st.dataframe(pd.DataFrame({
        'Date': [to_ddmmyyyy(day['date']) for day in st.session_state.exam_schedule],
        'First Half Required': [day['first_half'] for day in st.session_state.exam_schedule],
        'Second Half Required': [day['second_half'] for day in st.session_state.exam_schedule]
    }))
    st.subheader("Faculty List")
    st.dataframe(st.session_state.faculty_df)
    st.markdown("---")
    st.subheader("Assignment Generation")
    st.info("Assignments will respect all constraints: unavailability, max duties, groupings, and required per shift.")

    if st.button("🎲 Generate Assignments"):
        faculty_list = st.session_state.faculty_df['faculty'].tolist()
        max_duties_dict = st.session_state.max_duties_dict.copy()
        unavailability = st.session_state.faculty_unavailability
        faculty_groups = st.session_state.faculty_groups
        exam_schedule = st.session_state.exam_schedule

        # --- DIAGNOSTICS BEFORE SOLVER ---
        st.markdown('---')
        st.subheader('Assignment Feasibility Diagnostics')
        # 1. Total required duties
        total_required = sum(day['first_half'] + day['second_half'] for day in exam_schedule)
        st.markdown(f'- **Total required duties:** {total_required}')
        # 2. Total available duties
        total_available = sum(int(st.session_state.max_duties_dict.get(f, 0)) for f in faculty_list)
        st.markdown(f'- **Total available duties (sum of max duties):** {total_available}')
        # 3. Slot-by-slot feasibility
        slot_problems = False
        for day in exam_schedule:
            for shift, label in [('First Half', 'first_half'), ('Second Half', 'second_half')]:
                required = day[label]
                date = day['date']
                available_faculty = [
                    f for f in faculty_list
                    if date not in unavailability[f][label]
                    and int(st.session_state.max_duties_dict.get(f, 0)) > 0
                ]
                st.markdown(f'- {to_ddmmyyyy(date)} {shift}: required = {required}, available = {len(available_faculty)}')
                if required > len(available_faculty):
                    st.warning(f'⚠️ Not enough available faculty for {to_ddmmyyyy(date)} {shift}: required {required}, available {len(available_faculty)}')
                    slot_problems = True

        # 4. Per-day feasibility (no both-shifts constraint)
        day_problems = False
        for day in exam_schedule:
            date = day['date']
            required_total = day['first_half'] + day['second_half']
            available_faculty = [
                f for f in faculty_list
                if date not in unavailability[f]['first_half']
                and date not in unavailability[f]['second_half']
                and int(st.session_state.max_duties_dict.get(f, 0)) > 0
            ]
            st.markdown(f'- {to_ddmmyyyy(date)}: total required = {required_total}, available = {len(available_faculty)}')
            if required_total > len(available_faculty):
                st.error(f'❌ Not enough available faculty for {to_ddmmyyyy(date)}: required {required_total}, available {len(available_faculty)}')
                day_problems = True

        if total_required > total_available:
            st.error('❌ Total required duties exceed total available duties. Assignment is impossible.')
            st.stop()
        elif slot_problems or day_problems:
            st.error('❌ At least one slot or day has more required than available faculty. Assignment is impossible.')
            st.stop()
        else:
            st.success('✅ Feasibility check passed. Proceeding to assignment solver...')
        # --- END DIAGNOSTICS ---

        # --- Heuristic Assignment Algorithm (FET-like) ---
        def generate_assignments_heuristic(faculty_list, max_duties_dict, unavailability, faculty_groups, exam_schedule):
            import copy
            # Build slot list: (date, shift, required)
            slots = []
            for day in exam_schedule:
                for shift, label in [("First Half", "first_half"), ("Second Half", "second_half")]:
                    required = day[label]
                    if required > 0:
                        slots.append({"date": day["date"], "shift": shift, "required": required})
            # Build group lookup
            group_map = {}
            for group in faculty_groups:
                for f in group:
                    group_map[f] = tuple(sorted(group))
            # Helper: get all groups (as tuples)
            all_groups = [tuple(sorted(g)) for g in faculty_groups]
            # Helper: get all singletons (faculty not in any group)
            grouped = set(f for g in faculty_groups for f in g)
            singletons = [f for f in faculty_list if f not in grouped]
            # Build assignable units: groups + singletons
            assign_units = all_groups + [(f,) for f in singletons]
            # Helper: check if a unit is available for a slot
            def unit_available(unit, date, shift, duty_counts, assigned_slots):
                for f in unit:
                    label = 'first_half' if shift == 'First Half' else 'second_half'
                    if date in unavailability[f][label]:
                        return False
                    if duty_counts.get(f, 0) >= int(max_duties_dict.get(f, 0)):
                        return False
                    for s in ["First Half", "Second Half"]:
                        if s != shift and (date, s, f) in assigned_slots:
                            return False
                return True
            def assign_unit(unit, date, shift, duty_counts, assigned_slots):
                for f in unit:
                    duty_counts[f] = duty_counts.get(f, 0) + 1
                    assigned_slots.add((date, shift, f))
            def unassign_unit(unit, date, shift, duty_counts, assigned_slots):
                for f in unit:
                    duty_counts[f] -= 1
                    assigned_slots.remove((date, shift, f))
            def slot_difficulty(slot, duty_counts, assigned_slots):
                avail = [u for u in assign_units if unit_available(u, slot["date"], slot["shift"], duty_counts, assigned_slots)]
                return len(avail)
            # --- Diagnostics: Check per-slot feasibility before assignment ---
            diagnostics_msgs = []
            temp_duty_counts = {f: 0 for f in faculty_list}
            temp_assigned_slots = set()
            infeasible = False
            for slot in slots:
                avail = [u for u in assign_units if unit_available(u, slot["date"], slot["shift"], temp_duty_counts, temp_assigned_slots)]
                diagnostics_msgs.append(f"{slot['date']} {slot['shift']}: required = {slot['required']}, available units = {len(avail)}")
                if len(avail) < slot['required']:
                    diagnostics_msgs.append(f"❌ Not enough available faculty/groups for {slot['date']} {slot['shift']}: required {slot['required']}, available {len(avail)}")
                    infeasible = True
            if infeasible:
                for msg in diagnostics_msgs:
                    print(msg)
                return None
            # Main backtracking assignment
            def backtrack(slot_idx, slots, duty_counts, assigned_slots, assignment):
                if slot_idx == len(slots):
                    return True
                slot = slots[slot_idx]
                date, shift, required = slot["date"], slot["shift"], slot["required"]
                available_units = [u for u in assign_units if unit_available(u, date, shift, duty_counts, assigned_slots)]
                available_units = sorted(available_units, key=lambda u: sum(duty_counts.get(f, 0) for f in u))
                from itertools import combinations
                for units in combinations(available_units, required):
                    flat = [f for u in units for f in u]
                    if len(set(flat)) != len(flat):
                        continue
                    for u in units:
                        assign_unit(u, date, shift, duty_counts, assigned_slots)
                    assignment.append({"date": date, "shift": shift, "faculty": [f for u in units for f in u]})
                    if backtrack(slot_idx + 1, slots, duty_counts, assigned_slots, assignment):
                        return True
                    assignment.pop()
                    for u in units:
                        unassign_unit(u, date, shift, duty_counts, assigned_slots)
                return False
            slots_sorted = sorted(slots, key=lambda s: slot_difficulty(s, temp_duty_counts, temp_assigned_slots))
            duty_counts = {f: 0 for f in faculty_list}
            assigned_slots = set()
            assignment = []
            found = backtrack(0, slots_sorted, duty_counts, assigned_slots, assignment)
            if not found:
                # --- Greedy fallback: assign as many as possible to each slot ---
                print("Backtracking failed, trying greedy fallback...")
                duty_counts = {f: 0 for f in faculty_list}
                assigned_slots = set()
                assignment = []
                underfilled_slots = []
                for slot in slots_sorted:
                    date, shift, required = slot["date"], slot["shift"], slot["required"]
                    available_units = [u for u in assign_units if unit_available(u, date, shift, duty_counts, assigned_slots)]
                    available_units = sorted(available_units, key=lambda u: sum(duty_counts.get(f, 0) for f in u))
                    chosen = []
                    used = set()
                    for u in available_units:
                        if len(chosen) >= required:
                            break
                        if not any(f in used for f in u):
                            chosen.append(u)
                            used.update(u)
                    if len(chosen) < required:
                        print(f"[Greedy] Underfilled {date} {shift}: required {required}, assigned {len(chosen)}")
                        underfilled_slots.append({"date": date, "shift": shift, "required": required, "assigned": len(chosen)})
                    for u in chosen:
                        assign_unit(u, date, shift, duty_counts, assigned_slots)
                    assignment.append({"date": date, "shift": shift, "faculty": [f for u in chosen for f in u]})
                if not assignment:
                    print("Greedy fallback also failed: no assignments made.")
                    return None
                # Return both assignment and underfilled_slots for UI warning
                return pd.DataFrame([{"Date": slot["date"], "Shift": slot["shift"], "Faculty": f} for slot in assignment for f in slot["faculty"]]), underfilled_slots
            # Build assignment DataFrame
            rows = []
            for slot in assignment:
                for f in slot["faculty"]:
                    rows.append({"Date": slot["date"], "Shift": slot["shift"], "Faculty": f})
            return pd.DataFrame(rows), []
        # Use heuristic assignment if possible
        assignment, underfilled_slots = generate_assignments_heuristic(faculty_list, max_duties_dict, unavailability, faculty_groups, exam_schedule)
        if assignment is None:
            st.stop()
        
        assigned_df = assignment
        assigned_df = normalize_dates(assigned_df)
        st.session_state.assigned_duty_df = assigned_df
        # Prefill manual_selected for manual intervention
        if 'manual_selected' not in st.session_state:
            st.session_state.manual_selected = {}
        st.session_state.manual_selected.clear()
        for _, row in assigned_df.iterrows():
            slot_key = f"slot_{to_ddmmyyyy(row['Date'])}_{row['Shift']}"
            if slot_key not in st.session_state.manual_selected:
                st.session_state.manual_selected[slot_key] = []
            st.session_state.manual_selected[slot_key].append(row['Faculty'])
        # Add contact info if available
        if not assigned_df.empty:
            lookup = st.session_state.faculty_df.set_index('faculty').to_dict(orient='index')
            for col in ['Phone No', 'Email Id', 'Designation']:
                if col in st.session_state.faculty_df.columns:
                    assigned_df[col] = assigned_df['Faculty'].map(lambda f: lookup.get(f, {}).get(col, ''))
        st.success("Assignments generated with heuristic!")
        # UI warning for under-filled slots
        if underfilled_slots:
            st.warning("⚠️ Some slots could not be fully filled:")
            for slot in underfilled_slots:
                st.warning(f"{to_ddmmyyyy(slot['date'])} {slot['shift']}: assigned {slot['assigned']} / required {slot['required']}")
        # Constraint validation and warnings
        is_valid, errors, same_day_double = validate_assignment_constraints(
            assigned_df,
            faculty_list,
            st.session_state.max_duties_dict,
            st.session_state.faculty_unavailability,
            st.session_state.faculty_groups,
            st.session_state.exam_schedule
        )
        warning_msgs = []
        error_msgs = []
        for err in errors:
            if any(f in err for f, d in same_day_double):
                warning_msgs.append(err)
            else:
                error_msgs.append(err)
        if error_msgs:
            st.warning("⚠️ Assignment constraint violations detected:")
            for err in error_msgs:
                st.error(err)
        if warning_msgs:
            st.warning("⚠️ The following faculty are assigned to both shifts on the same day (allowed with warning):")
            for warn in warning_msgs:
                st.warning(warn)
        if error_msgs:
            if st.button("Continue with Violations"):
                st.session_state.assigned_duty_df = assigned_df
                st.success("Assignments updated with violations!")
                # Regenerate reports immediately after update
                df = assigned_df.copy()
                excel_data = generate_faculty_summary_excel(df, st.session_state.faculty_unavailability)
                word_data = generate_word_doc(df)
            else:
                excel_data = generate_faculty_summary_excel(assigned_df.copy(), st.session_state.faculty_unavailability)
                word_data = generate_word_doc(assigned_df.copy())
        else:
            st.session_state.assigned_duty_df = assigned_df
            st.success("Assignments updated!")
            # Regenerate reports immediately after update
            df = assigned_df.copy()
            excel_data = generate_faculty_summary_excel(df, st.session_state.faculty_unavailability)
            word_data = generate_word_doc(df)
        st.dataframe(assigned_df)
        # Add report regeneration/download buttons
        st.markdown("---")
        st.markdown("### Regenerate & Download Reports")
        if 'excel_data' not in locals():
            excel_data = generate_faculty_summary_excel(df.copy(), st.session_state.faculty_unavailability)
        if 'word_data' not in locals():
            word_data = generate_word_doc(df.copy())
        # Show status of report generation
        if excel_data is not None:
            st.success("✅ Excel report generated successfully")
        else:
            st.error("❌ Failed to generate Excel report")
        if word_data is not None:
            st.success("✅ Word report generated successfully")
        else:
            st.error("❌ Failed to generate Word report")
        col1, col2 = st.columns(2)
        with col1:
            if excel_data is not None:
                st.download_button(
                    "📊 Download Excel Report",
                    excel_data,
                    "faculty_summary.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        with col2:
            if word_data is not None:
                st.download_button(
                    "📄 Download Word Report",
                    word_data,
                    "faculty_duty_assignment.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # --- Shift duty warning ---
        max_duties_dict = st.session_state.max_duties_dict
        faculty_list = st.session_state.faculty_df['faculty'].astype(str).str.strip().tolist()
        for faculty in faculty_list:
            max_duty = int(max_duties_dict.get(faculty, 0))
            fh_count = len(assigned_df[(assigned_df['Faculty'] == faculty) & (assigned_df['Shift'] == 'First Half')])
            sh_count = len(assigned_df[(assigned_df['Faculty'] == faculty) & (assigned_df['Shift'] == 'Second Half')])
            limit = (max_duty + 1) // 2
            if fh_count > limit:
                st.warning(f"⚠️ {faculty} has {fh_count} First Half duties, which exceeds (max duty + 1) // 2 = {limit}.")
            if sh_count > limit:
                st.warning(f"⚠️ {faculty} has {sh_count} Second Half duties, which exceeds (max duty + 1) // 2 = {limit}.")

# --- Manual Intervention ---
elif section == "Manual Intervention":
    st.header("Manual Intervention")
    # At the start of Manual Intervention section
    if 'manual_selected' not in st.session_state:
        st.session_state.manual_selected = {}
    # Expander for uploading and regenerating from Excel summary
    with st.expander("Upload and Regenerate from Excel Summary", expanded=False):
        uploaded_summary = st.file_uploader("Upload Excel Summary (faculty_summary.xlsx)", type=["xlsx"], key="manual_excel_upload")
        if uploaded_summary is not None:
            try:
                summary_df = pd.read_excel(uploaded_summary)
                # Try to reconstruct assignment DataFrame from summary (if possible)
                # This expects columns: Faculty, First Half Dates, Second Half Dates
                assignment_rows = []
                for _, row in summary_df.iterrows():
                    faculty = row.get('Faculty', '')
                    for col, shift in [("First Half Dates", "First Half"), ("Second Half Dates", "Second Half")]:
                        if pd.notna(row.get(col, '')) and str(row.get(col, '')).strip():
                            date_strs = [d.strip() for d in str(row[col]).split(',') if d.strip()]
                            for date_str in date_strs:
                                date = from_ddmmyyyy(date_str)
                                if date:
                                    assignment_rows.append({'Date': date, 'Shift': shift, 'Faculty': faculty})
                new_df = pd.DataFrame(assignment_rows)
                # Add contact info if available
                if not new_df.empty and st.session_state.faculty_df is not None:
                    lookup = st.session_state.faculty_df.set_index('faculty').to_dict(orient='index')
                    for col in ['Phone No', 'Email Id', 'Designation']:
                        if col in st.session_state.faculty_df.columns:
                            new_df[col] = new_df['Faculty'].map(lambda f: lookup.get(f, {}).get(col, ''))
                st.session_state.assigned_duty_df = new_df
                # Prefill manual_selected for manual intervention
                if 'manual_selected' not in st.session_state:
                    st.session_state.manual_selected = {}
                st.session_state.manual_selected.clear()
                for _, row in new_df.iterrows():
                    slot_key = f"slot_{to_ddmmyyyy(row['Date'])}_{row['Shift']}"
                    if slot_key not in st.session_state.manual_selected:
                        st.session_state.manual_selected[slot_key] = []
                    st.session_state.manual_selected[slot_key].append(row['Faculty'])
                st.success("Assignments regenerated from uploaded summary!")
                st.dataframe(new_df)
                # Regenerate Excel/Word reports immediately
                excel_data = generate_faculty_summary_excel(new_df, st.session_state.faculty_unavailability)
                word_data = generate_word_doc(new_df)
            except Exception as e:
                st.error(f"Failed to parse uploaded summary: {e}")
        auto_save_job()
    if st.session_state.assigned_duty_df is not None and not st.session_state.assigned_duty_df.empty:
        df = st.session_state.assigned_duty_df.copy()
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
        # Defensive check for required columns
        required_cols = ['Date', 'Shift', 'Faculty']
        for col in required_cols:
            if col not in df.columns:
                st.error(f"Assignment data is missing required column: '{col}'. Please check your input.")
                st.stop()
        if df.empty:
            df = pd.DataFrame(columns=required_cols)
        if st.session_state.faculty_df is not None:
            faculty_list = st.session_state.faculty_df['faculty'].astype(str).str.strip().tolist()
            # Build a unique list of (date, shift) slots
            slots = df[['Date', 'Shift']].drop_duplicates().sort_values(['Date', 'Shift']).values.tolist()
            new_assignments = []
            st.markdown('#### Edit Assignments by Slot')
            # Ensure unavailability and max_duties_dict are defined
            unavailability = st.session_state.faculty_unavailability
            max_duties_dict = st.session_state.max_duties_dict
            # Recalculate current assignments for all slots after every change
            # Build a working assignments dict from session_state.manual_selected
            working_assignments = {}
            for s_date, s_shift in slots:
                s_slot_key = f"slot_{to_ddmmyyyy(s_date)}_{s_shift}"
                working_assignments[(s_date, s_shift)] = st.session_state.manual_selected.get(s_slot_key, [])
            for date, shift in slots:
                slot_df = df[(df['Date'] == date) & (df['Shift'] == shift)]
                assigned_faculty = slot_df['Faculty'].astype(str).str.strip().tolist()
                required = 0
                for day in st.session_state.exam_schedule:
                    if day['date'] == date:
                        required = day['first_half'] if shift == "First Half" else day['second_half']
                        break
                assigned_on_date = df[df['Date'] == date]['Faculty'].astype(str).str.strip().tolist()
                assigned_on_slot = [f for f in assigned_faculty if f]  # filter out empty strings
                st.markdown(f"**{to_ddmmyyyy(date)} - {shift}**")
                slot_key = f"slot_{to_ddmmyyyy(date)}_{shift}"
                # Make a copy of working_assignments and update with the current selection for this slot
                temp_assignments = working_assignments.copy()
                temp_assignments[(date, shift)] = st.session_state.manual_selected.get(slot_key, [])
                # Now calculate counts from temp_assignments
                all_assigned = []
                for (s_date, s_shift), facs in temp_assignments.items():
                    all_assigned.extend(facs)
                current_duty_counts = pd.Series(all_assigned).value_counts().to_dict()
                fh_counts = pd.Series([f for (d, s), facs in temp_assignments.items() if s == 'First Half' for f in facs]).value_counts().to_dict()
                sh_counts = pd.Series([f for (d, s), facs in temp_assignments.items() if s == 'Second Half' for f in facs]).value_counts().to_dict()
                # For 'no both shifts' rule, get faculty assigned to the other shift on this date
                other_shift = 'Second Half' if shift == 'First Half' else 'First Half'
                assigned_to_other_shift = temp_assignments.get((date, other_shift), [])
                label = 'first_half' if shift == 'First Half' else 'second_half'
                # Build available_faculty: all eligible + all already selected for this slot
                eligible_faculty = [
                    f for f in faculty_list
                    if (
                        (
                            date not in unavailability[f][label] and
                            f not in assigned_to_other_shift and
                            current_duty_counts.get(f, 0) < int(max_duties_dict.get(f, 0))
                        )
                    )
                ]
                # Always include already selected faculty for this slot (even if over max)
                available_faculty = sorted(set(eligible_faculty) | set(assigned_on_slot))
                # Build display options and mapping
                display_to_faculty = {}
                faculty_to_display = {}
                display_options = []
                for f in available_faculty:
                    total = current_duty_counts.get(f, 0)
                    fh = fh_counts.get(f, 0)
                    sh = sh_counts.get(f, 0)
                    display = f"{f} (Total: {total}, First Half: {fh}, Second Half: {sh})"
                    display_options.append(display)
                    display_to_faculty[display] = f
                    faculty_to_display[f] = display
                # Now define valid_assigned_on_slot after faculty_to_display is built
                valid_assigned_on_slot = [f for f in assigned_on_slot if f in faculty_to_display]
                # Trim to max allowed
                if required > 0 and len(valid_assigned_on_slot) > required:
                    valid_assigned_on_slot = valid_assigned_on_slot[:required]
                if slot_key not in st.session_state.manual_selected:
                    st.session_state.manual_selected[slot_key] = valid_assigned_on_slot
                # Also filter session_state.manual_selected[slot_key] to valid options and max_selections
                st.session_state.manual_selected[slot_key] = [
                    f for f in st.session_state.manual_selected[slot_key] if f in faculty_to_display
                ][:required if required > 0 else None]
                default_display = [faculty_to_display[f] for f in st.session_state.manual_selected[slot_key] if f in faculty_to_display]
                selected_display = st.multiselect(
                    f"Assign Faculty for {to_ddmmyyyy(date)} {shift} (Required: {required})",
                    options=display_options,
                    default=default_display,
                    key=slot_key,
                    max_selections=required if required > 0 else None
                )
                # Map back to faculty names for saving and display
                selected = [display_to_faculty[d] for d in selected_display]
                st.session_state.manual_selected[slot_key] = selected
                assigned_count = len(selected)
                color = '#5cb85c' if assigned_count == required else '#d9534f'
                st.markdown(f"<span style='color: {color}; font-weight: 600;'>{assigned_count} faculty assigned on {to_ddmmyyyy(date)} {shift}, required: {required}.</span>", unsafe_allow_html=True)
                if selected:
                    st.markdown("**Selected Faculty:**<br>" + ", ".join(selected), unsafe_allow_html=True)
                for f in selected:
                    new_assignments.append({'Date': date, 'Shift': shift, 'Faculty': f})
            new_df = pd.DataFrame(new_assignments)
            # Defensive check for required columns
            for col in required_cols:
                if col not in new_df.columns:
                    new_df[col] = None
            if not new_df.empty:
                lookup = st.session_state.faculty_df.set_index('faculty').to_dict(orient='index')
                for col in ['Phone No', 'Email Id', 'Designation']:
                    if col in st.session_state.faculty_df.columns:
                        new_df[col] = new_df['Faculty'].map(lambda f: lookup.get(f, {}).get(col, ''))
            is_valid, errors, same_day_double = validate_assignment_constraints(
                new_df,
                faculty_list,
                st.session_state.max_duties_dict,
                st.session_state.faculty_unavailability,
                st.session_state.faculty_groups,
                st.session_state.exam_schedule
            )
            # In Manual Intervention, show both-shift errors as warnings, others as errors
            warning_msgs = []
            error_msgs = []
            for err in errors:
                if any(f in err for f, d in same_day_double):
                    warning_msgs.append(err)
                else:
                    error_msgs.append(err)
            if error_msgs:
                st.warning("⚠️ Assignment constraint violations detected:")
                for err in error_msgs:
                    st.error(err)
            if warning_msgs:
                st.warning("⚠️ The following faculty are assigned to both shifts on the same day (allowed with warning):")
                for warn in warning_msgs:
                    st.warning(warn)
            if error_msgs:
                if st.button("Continue with Violations"):
                    st.session_state.assigned_duty_df = new_df
                    st.success("Assignments updated with violations!")
                    # Regenerate reports immediately after update
                    df = new_df.copy()
                    excel_data = generate_faculty_summary_excel(df, st.session_state.faculty_unavailability)
                    word_data = generate_word_doc(df)
                else:
                    excel_data = generate_faculty_summary_excel(df.copy(), st.session_state.faculty_unavailability)
                    word_data = generate_word_doc(df.copy())
            else:
                st.session_state.assigned_duty_df = new_df
                st.success("Assignments updated!")
                # Regenerate reports immediately after update
                df = new_df.copy()
                excel_data = generate_faculty_summary_excel(df, st.session_state.faculty_unavailability)
                word_data = generate_word_doc(df)
            st.dataframe(new_df)
            # Add report regeneration/download buttons
            st.markdown("---")
            st.markdown("### Regenerate & Download Reports")
            if 'excel_data' not in locals():
                excel_data = generate_faculty_summary_excel(df.copy(), st.session_state.faculty_unavailability)
            if 'word_data' not in locals():
                word_data = generate_word_doc(df.copy())
            # Show status of report generation
            if excel_data is not None:
                st.success("✅ Excel report generated successfully")
            else:
                st.error("❌ Failed to generate Excel report")
            if word_data is not None:
                st.success("✅ Word report generated successfully")
            else:
                st.error("❌ Failed to generate Word report")
            col1, col2 = st.columns(2)
            with col1:
                if excel_data is not None:
                    st.download_button(
                        "📊 Download Excel Report",
                        excel_data,
                        "faculty_summary.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            with col2:
                if word_data is not None:
                    st.download_button(
                        "📄 Download Word Report",
                        word_data,
                        "faculty_duty_assignment.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            # --- Shift duty warning ---
            max_duties_dict = st.session_state.max_duties_dict
            faculty_list = st.session_state.faculty_df['faculty'].astype(str).str.strip().tolist()
            for faculty in faculty_list:
                max_duty = int(max_duties_dict.get(faculty, 0))
                fh_count = len(new_df[(new_df['Faculty'] == faculty) & (new_df['Shift'] == 'First Half')])
                sh_count = len(new_df[(new_df['Faculty'] == faculty) & (new_df['Shift'] == 'Second Half')])
                limit = (max_duty + 1) // 2
                if fh_count > limit:
                    st.warning(f"⚠️ {faculty} has {fh_count} First Half duties, which exceeds (max duty + 1) // 2 = {limit}.")
                if sh_count > limit:
                    st.warning(f"⚠️ {faculty} has {sh_count} Second Half duties, which exceeds (max duty + 1) // 2 = {limit}.")
        else:
            st.warning("Faculty list is not loaded. Please upload faculty details first.")
            st.stop()
    else:
        st.warning("No assignments available to edit. You can start manual assignment from scratch.")
        if st.button("Start Manual Assignment"):
            # Build empty DataFrame with all slots
            slots = []
            for day in st.session_state.exam_schedule:
                for shift in ["First Half", "Second Half"]:
                    required = day['first_half'] if shift == "First Half" else day['second_half']
                    for _ in range(required):
                        slots.append({'Date': day['date'], 'Shift': shift, 'Faculty': ''})
            new_df = pd.DataFrame(slots)
            st.session_state.assigned_duty_df = new_df
            st.success("Manual assignment initialized. Please assign faculty to each slot.")
            st.rerun()
        auto_save_job()

# --- Reports & Downloads ---
elif section == "Reports & Downloads":
    st.header("Reports & Downloads")
    if st.session_state.assigned_duty_df is not None and not st.session_state.assigned_duty_df.empty:
        df = st.session_state.assigned_duty_df.copy()
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
        # Format for display
        df_display = df.copy()
        if 'Date' in df_display.columns:
            df_display['Date'] = df_display['Date'].apply(to_ddmmyyyy)
        df_display.insert(0, 'S.No', range(1, len(df_display) + 1))
        st.dataframe(df_display)
        # Download buttons (Excel, Word, CSV) using gd45.py logic
        excel_data = generate_faculty_summary_excel(
            df.copy(),
            st.session_state.faculty_unavailability
        )
        word_data = generate_word_doc(df.copy())
        col1, col2 = st.columns(2)
        with col1:
            if excel_data is not None:
                st.download_button(
                    "📊 Download Excel Report",
                    excel_data,
                    "faculty_summary.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        with col2:
            if word_data is not None:
                st.download_button(
                    "📄 Download Word Report",
                    word_data,
                    "faculty_duty_assignment.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("No assignments available to download.")

# --- End of App ---

st.markdown(
    '''
    <style>
    div[data-baseweb="select"] {
        min-width: 400px !important;
        max-width: 100% !important;
    }
    </style>
    ''',
    unsafe_allow_html=True
) 
